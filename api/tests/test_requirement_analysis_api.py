import io
from unittest.mock import AsyncMock, patch

import pytest
from docx import Document
from fastapi.testclient import TestClient

from services.database import ensure_initial_admin, init_db


@pytest.fixture(autouse=True)
def temp_db(tmp_path, monkeypatch):
    db_path = str(tmp_path / "test_requirement_analysis.db")
    monkeypatch.setattr("services.database.get_db_path", lambda: db_path)
    monkeypatch.setattr("services.production_issue_file_store.get_db_path", lambda: db_path)
    monkeypatch.setattr("services.test_issue_file_store.get_db_path", lambda: db_path)
    monkeypatch.setenv("SESSION_SECRET", "test-session-secret")
    monkeypatch.setenv("INITIAL_ADMIN_USERNAME", "admin")
    monkeypatch.setenv("INITIAL_ADMIN_PASSWORD", "password123")
    monkeypatch.setenv("INITIAL_ADMIN_DISPLAY_NAME", "测试管理员")
    init_db()
    ensure_initial_admin()
    return db_path


@pytest.fixture
def client():
    from index import app

    return TestClient(app)


def build_requirement_docx() -> bytes:
    document = Document()
    document.add_paragraph("4.1 功能描述")
    document.add_paragraph("回家活动需要补充资格校验，并新增投保页面，避免历史缺陷中的资格校验遗漏。")
    document.add_paragraph("4.4 界面")
    document.add_paragraph("回家活动页面需要展示资格提示，并补充弹窗内容核对。")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_legacy_word_bytes() -> bytes:
    return bytes.fromhex("D0CF11E0A1B11AE1") + (b"\x00" * 512)


def prepare_analysis_context(client: TestClient) -> int:
    login_response = client.post(
        "/api/auth/login",
        json={"username": "admin", "password": "password123"},
    )
    assert login_response.status_code == 200

    project_response = client.post("/api/projects", json={"name": "需求分析项目", "description": ""})
    assert project_response.status_code == 200
    project_id = project_response.json()["data"]["id"]

    mapping_response = client.put(
        f"/api/projects/{project_id}/requirement-mapping",
        json={
            "groups": [
                {
                    "id": "group-1",
                    "tag": "页面新增",
                    "requirement_keyword": "新增页面",
                    "related_scenarios": ["兼容性测试", "跳转链路"],
                },
                {
                    "id": "group-2",
                    "tag": "弹窗",
                    "requirement_keyword": "新增弹窗",
                    "related_scenarios": ["弹窗内容核对", "弹窗页面其他弹窗相关性测试"],
                },
            ]
        },
    )
    assert mapping_response.status_code == 200
    return project_id


def test_requirement_analysis_creates_record_and_supports_history(client: TestClient):
    project_id = prepare_analysis_context(client)

    response = client.post(
        "/api/requirement-analysis/analyze",
        data={
            "project_id": str(project_id),
            "use_ai": "false",
        },
        files={
            "requirement_file": (
                "requirement.docx",
                build_requirement_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert response.status_code == 200
    payload = response.json()["data"]
    assert payload["overview"]["use_ai"] is False
    assert payload["overview"]["matched_requirements"] >= 1
    assert payload["mapping_suggestions"]
    assert payload["overview"]["mapping_hit_count"] >= 1
    assert payload["source_files"]["requirement_mapping_available"] is True
    assert "production_alerts" not in payload
    assert "test_suggestions" not in payload
    assert payload["record_id"] > 0

    records_response = client.get("/api/requirement-analysis/records")
    assert records_response.status_code == 200
    records = records_response.json()["data"]
    assert len(records) == 1
    assert records[0]["requirement_file_name"] == "requirement.docx"
    assert records[0]["mapping_hit_count"] >= 1

    detail_response = client.get(f"/api/requirement-analysis/records/{payload['record_id']}")
    assert detail_response.status_code == 200
    detail = detail_response.json()["data"]
    assert detail["section_snapshot"]["selected_mode"] == "preferred_sections"
    assert detail["result_snapshot"]["overview"]["matched_requirements"] >= 1
    assert detail["result_snapshot"]["mapping_suggestions"]


def test_requirement_analysis_accepts_doc_files(client: TestClient):
    project_id = prepare_analysis_context(client)

    with patch(
        "index.parse_requirement_document",
        return_value={
            "selected_mode": "preferred_sections",
            "document_type": "doc",
            "selected_sections": [
                {"number": "4.1", "title": "功能描述", "block_count": 1},
                {"number": "4.4", "title": "界面", "block_count": 1},
            ],
            "all_sections": [
                {"number": "4.1", "title": "功能描述", "block_count": 1},
                {"number": "4.4", "title": "界面", "block_count": 1},
            ],
            "points": [
                {
                    "point_id": "4.1-1",
                    "section_number": "4.1",
                    "section_title": "功能描述",
                    "text": "山东济宁中支明白纸优化，增加温馨提示语音播报。",
                }
            ],
        },
    ):
        response = client.post(
            "/api/requirement-analysis/analyze",
            data={
                "project_id": str(project_id),
                "use_ai": "false",
            },
            files={
                "requirement_file": (
                    "legacy.doc",
                    build_legacy_word_bytes(),
                    "application/msword",
                )
            },
        )

    assert response.status_code == 200
    assert response.json()["data"]["source_files"]["requirement_file_name"] == "legacy.doc"


def test_requirement_analysis_uses_deepseek_when_enabled(client: TestClient):
    project_id = prepare_analysis_context(client)

    with patch(
        "index.parse_requirement_document",
        return_value={
            "selected_mode": "preferred_sections",
            "document_type": "docx",
            "selected_sections": [
                {"number": "4.1", "title": "功能描述", "block_count": 1},
                {"number": "4.4", "title": "界面", "block_count": 1},
            ],
            "all_sections": [
                {"number": "4.1", "title": "功能描述", "block_count": 1},
                {"number": "4.4", "title": "界面", "block_count": 1},
            ],
            "points": [
                {
                    "point_id": "4.1-1",
                    "section_number": "4.1",
                    "section_title": "功能描述",
                    "text": "本次需要补充新增页面相关验证。",
                }
            ],
        },
    ), patch(
        "index.call_deepseek",
        AsyncMock(
            return_value={
                "result": {
                    "summary": "新增页面需求命中需求映射关系，建议补齐同组关联场景回归。",
                    "overall_assessment": "映射扩展场景需要重点回归，优先校验页面兼容性与跳转链路。",
                    "key_findings": [
                        "新增页面命中后，需要补齐同组关联场景。",
                        "新增页面命中后，需要补齐同组关联场景。",
                        "建议优先验证兼容性测试和跳转链路。",
                    ],
                    "risk_table": [
                        {
                            "requirement_point_id": "4.1-1",
                            "risk_level": "高",
                            "risk_reason": "命中需求映射后，需要扩展到整组关联场景。",
                            "test_focus": "优先验证兼容性测试和跳转链路。",
                        },
                        {
                            "requirement_point_id": "4.1-1",
                            "risk_level": "高",
                            "risk_reason": "重复行",
                            "test_focus": "重复行",
                        },
                    ],
                },
                "usage": {
                    "prompt_tokens": 100,
                    "completion_tokens": 60,
                    "total_tokens": 160,
                    "prompt_cache_hit_tokens": 0,
                    "prompt_cache_miss_tokens": 100,
                },
            }
        ),
    ):
        response = client.post(
            "/api/requirement-analysis/analyze",
            data={
                "project_id": str(project_id),
                "use_ai": "true",
            },
            files={
                "requirement_file": (
                    "requirement.docx",
                    build_requirement_docx(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    assert response.status_code == 200
    payload = response.json()["data"]
    assert payload["ai_analysis"]["provider"] == "DeepSeek"
    assert payload["ai_analysis"]["summary"] == "新增页面需求命中需求映射关系，建议补齐同组关联场景回归。"
    assert payload["ai_analysis"]["overall_assessment"] == "映射扩展场景需要重点回归"
    assert payload["ai_analysis"]["key_findings"] == [
        "新增页面命中后，需要补齐同组关联场景。",
        "建议优先验证兼容性测试和跳转链路。",
    ]
    assert len(payload["ai_analysis"]["risk_table"]) == 1
    assert payload["ai_analysis"]["risk_table"][0]["risk_level"] == "高"
    assert payload["ai_cost"]["total_tokens"] == 160
    assert payload["mapping_suggestions"][0]["suggestion"]
    assert payload["source_files"]["requirement_mapping_available"] is True
    assert "production_alerts" not in payload
    assert "test_suggestions" not in payload
