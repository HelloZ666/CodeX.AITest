import io
import json
from unittest.mock import AsyncMock, patch

import pytest
from docx import Document
from fastapi.testclient import TestClient

from services.database import ensure_initial_admin, init_db


@pytest.fixture(autouse=True)
def temp_db(tmp_path, monkeypatch):
    db_path = str(tmp_path / "test_requirement_case_generation.db")
    monkeypatch.setattr("services.database.get_db_path", lambda: db_path)
    monkeypatch.setattr("services.production_issue_file_store.get_db_path", lambda: db_path)
    monkeypatch.setattr("services.test_issue_file_store.get_db_path", lambda: db_path)
    monkeypatch.setattr("services.config_library_store.get_db_path", lambda: db_path)
    monkeypatch.setenv("SESSION_SECRET", "test-session-secret")
    monkeypatch.setenv("INITIAL_ADMIN_USERNAME", "admin")
    monkeypatch.setenv("INITIAL_ADMIN_PASSWORD", "password123")
    monkeypatch.setenv("INITIAL_ADMIN_DISPLAY_NAME", "测试管理员")
    init_db()
    ensure_initial_admin()
    return db_path


@pytest.fixture
def client():
    from index import app

    with TestClient(app) as test_client:
        login_resp = test_client.post(
            "/api/auth/login",
            json={"username": "admin", "password": "password123"},
        )
        assert login_resp.status_code == 200
        yield test_client


def build_requirement_docx() -> bytes:
    document = Document()
    document.add_paragraph("4.1 功能描述")
    document.add_paragraph("新增投保资格校验，未满足资格条件时禁止提交，并提示失败原因。")
    document.add_paragraph("4.4 界面")
    document.add_paragraph("资格校验失败时，页面需要展示明显提示文案和弹窗引导。")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def create_project(client: TestClient, name: str = "功能测试项目") -> dict:
    response = client.post("/api/projects", json={"name": name, "description": ""})
    assert response.status_code == 200
    return response.json()["data"]


def test_requirement_case_mapping_preview_does_not_persist(client: TestClient):
    project = create_project(client, "案例映射预览项目")

    response = client.post(
        "/api/functional-testing/case-generation/map",
        data={"project_id": str(project["id"])},
        files={
            "requirement_file": (
                "requirement.docx",
                build_requirement_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )

    assert response.status_code == 200
    payload = response.json()["data"]
    assert payload["overview"]["total_requirements"] >= 1
    assert payload["overview"]["matched_requirements"] == 0
    assert payload["overview"]["mapping_hit_count"] == 0
    assert payload["overview"]["unmatched_requirements"] >= 1
    assert payload["mapping_suggestions"] == []
    assert payload["requirement_hits"] == []
    assert payload["unmatched_requirements"]
    assert payload["ai_analysis"] is None
    assert payload["ai_cost"] is None
    assert payload["source_files"]["project_id"] == project["id"]
    assert payload["source_files"]["project_name"] == project["name"]
    assert payload["source_files"]["requirement_mapping_available"] is False

    requirement_records_resp = client.get("/api/requirement-analysis/records")
    assert requirement_records_resp.status_code == 200
    assert requirement_records_resp.json()["data"] == []

    requirement_docs_resp = client.get("/api/config-management/requirement-documents")
    assert requirement_docs_resp.status_code == 200
    assert requirement_docs_resp.json()["data"] == []


def test_requirement_case_generation_preview_does_not_persist(client: TestClient):
    project = create_project(client, "案例生成预览项目")
    mapping_snapshot = {
        "overview": {
            "total_requirements": 2,
            "matched_requirements": 0,
            "mapping_hit_count": 0,
            "unmatched_requirements": 2,
        },
        "mapping_suggestions": [],
        "requirement_hits": [],
        "unmatched_requirements": [
            {
                "point_id": "P1",
                "section_number": "4.1",
                "section_title": "功能描述",
                "text": "新增投保资格校验。",
            }
        ],
        "ai_analysis": None,
        "ai_cost": None,
    }

    with patch(
        "services.requirement_case_generator.call_deepseek",
        new=AsyncMock(
            return_value={
                "result": {
                    "summary": "覆盖资格校验主流程、失败拦截和界面提示。",
                    "cases": [
                        {
                            "case_id": "TC-001",
                            "description": "资格校验失败时禁止提交",
                            "steps": "1. 输入不满足条件的数据\n2. 提交",
                            "expected_result": "系统阻止提交并提示失败原因",
                        }
                    ],
                },
                "usage": {
                    "prompt_tokens": 120,
                    "completion_tokens": 80,
                    "total_tokens": 200,
                    "prompt_cache_hit_tokens": 0,
                    "prompt_cache_miss_tokens": 120,
                },
                "provider": "DeepSeek",
                "provider_key": "deepseek",
            }
        ),
    ) as mock_call:
        response = client.post(
            "/api/functional-testing/case-generation/generate",
            data={
                "project_id": str(project["id"]),
                "prompt_template_key": "requirement",
                "mapping_result_snapshot": json.dumps(mapping_snapshot, ensure_ascii=False),
                "reasoning_level": "high",
            },
            files={
                "requirement_file": (
                    "requirement.docx",
                    build_requirement_docx(),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            },
        )

    assert response.status_code == 200
    payload = response.json()["data"]
    assert payload["summary"] == "覆盖资格校验主流程、失败拦截和界面提示。"
    assert payload["generation_mode"] == "ai"
    assert payload["provider"] == "DeepSeek"
    assert payload["ai_cost"]["total_tokens"] == 200
    assert payload["total"] == 1
    assert payload["cases"][0]["case_id"] == "TC-001"
    assert "record_id" not in payload
    assert mock_call.await_args.kwargs["reasoning_level"] == "high"

    case_records_resp = client.get("/api/functional-testing/test-cases")
    assert case_records_resp.status_code == 200
    assert case_records_resp.json()["data"] == []

    requirement_docs_resp = client.get("/api/config-management/requirement-documents")
    assert requirement_docs_resp.status_code == 200
    assert requirement_docs_resp.json()["data"] == []

    case_assets_resp = client.get("/api/config-management/test-cases")
    assert case_assets_resp.status_code == 200
    assert case_assets_resp.json()["data"] == []

    audit_logs_resp = client.get("/api/audit-logs")
    assert audit_logs_resp.status_code == 200
    assert not any(
        item.get("request_path") == "/api/functional-testing/case-generation/generate"
        for item in audit_logs_resp.json()["data"]
    )


def test_requirement_case_save_persists_snapshots_without_ai_call(client: TestClient):
    project = create_project(client, "案例保存项目")
    requirement_bytes = build_requirement_docx()

    map_resp = client.post(
        "/api/functional-testing/case-generation/map",
        data={"project_id": str(project["id"])},
        files={
            "requirement_file": (
                "requirement.docx",
                requirement_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )
    assert map_resp.status_code == 200
    mapping_snapshot = map_resp.json()["data"]

    generation_snapshot = {
        "summary": "基于快照保存测试案例。",
        "generation_mode": "fallback",
        "provider": "DeepSeek",
        "ai_cost": None,
        "error": None,
        "total": 1,
        "cases": [
            {
                "case_id": "TC-001",
                "description": "资格校验失败时禁止提交",
                "steps": "1. 输入不满足条件的数据\n2. 提交",
                "expected_result": "系统阻止提交并提示失败原因",
                "source": "fallback",
            }
        ],
    }

    with patch(
        "services.requirement_case_generator.call_deepseek",
        new=AsyncMock(side_effect=AssertionError("save should not call ai")),
    ):
        save_resp = client.post(
            "/api/functional-testing/case-generation/save",
            data={
                "project_id": str(project["id"]),
                "prompt_template_key": "requirement",
                "case_name": "资格校验用例集",
                "iteration_version": "v1.2.0",
                "mapping_result_snapshot": json.dumps(mapping_snapshot, ensure_ascii=False),
                "generation_result_snapshot": json.dumps(generation_snapshot, ensure_ascii=False),
            },
            files={
                "requirement_file": (
                    "requirement.docx",
                    requirement_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
            },
        )

    assert save_resp.status_code == 200
    saved_payload = save_resp.json()["data"]
    assert isinstance(saved_payload["record_id"], int)
    assert saved_payload["name"] == "资格校验用例集"
    assert saved_payload["iteration_version"] == "v1.2.0"
    assert saved_payload["case_count"] == 1

    case_records_resp = client.get("/api/functional-testing/test-cases")
    assert case_records_resp.status_code == 200
    records = case_records_resp.json()["data"]
    assert len(records) == 1
    assert records[0]["id"] == saved_payload["record_id"]

    case_detail_resp = client.get(f"/api/functional-testing/test-cases/{saved_payload['record_id']}")
    assert case_detail_resp.status_code == 200
    detail = case_detail_resp.json()["data"]
    assert detail["name"] == "资格校验用例集"
    assert detail["iteration_version"] == "v1.2.0"
    assert detail["cases"][0]["case_id"] == "TC-001"

    requirement_records_resp = client.get("/api/requirement-analysis/records")
    assert requirement_records_resp.status_code == 200
    requirement_records = requirement_records_resp.json()["data"]
    assert len(requirement_records) == 1
    assert requirement_records[0]["requirement_file_name"] == "requirement.docx"

    requirement_docs_resp = client.get("/api/config-management/requirement-documents")
    assert requirement_docs_resp.status_code == 200
    requirement_docs = requirement_docs_resp.json()["data"]
    assert len(requirement_docs) == 1
    assert requirement_docs[0]["file_name"] == "requirement.docx"

    case_assets_resp = client.get("/api/config-management/test-cases")
    assert case_assets_resp.status_code == 200
    assets = case_assets_resp.json()["data"]
    assert len(assets) == 1
    assert assets[0]["name"] == "资格校验用例集"
    assert assets[0]["iteration_version"] == "v1.2.0"
    assert assets[0]["case_count"] == 1

    audit_logs_resp = client.get("/api/audit-logs")
    assert audit_logs_resp.status_code == 200
    save_logs = [
        item for item in audit_logs_resp.json()["data"]
        if item.get("request_path") == "/api/functional-testing/case-generation/save"
    ]
    assert len(save_logs) == 1
    assert save_logs[0]["action"] == "保存测试案例"
    assert save_logs[0]["detail"] == "已保存 1 条测试用例"
