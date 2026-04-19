import hashlib
import io
import json
from pathlib import Path
from unittest.mock import AsyncMock, patch

import pytest
from docx import Document
from fastapi.testclient import TestClient

from services.config_library_store import upsert_requirement_document, upsert_test_case_asset
from services.database import create_project as create_project_record, ensure_initial_admin, init_db


FIXTURES_DIR = Path(__file__).resolve().parent / "fixtures"


@pytest.fixture(autouse=True)
def temp_db(tmp_path, monkeypatch):
    db_path = str(tmp_path / "test_config_management_assets.db")
    monkeypatch.setattr("services.database.get_db_path", lambda: db_path)
    monkeypatch.setattr("services.production_issue_file_store.get_db_path", lambda: db_path)
    monkeypatch.setattr("services.test_issue_file_store.get_db_path", lambda: db_path)
    monkeypatch.setattr("services.config_library_store.get_db_path", lambda: db_path)
    monkeypatch.setenv("SESSION_SECRET", "test-session-secret")
    monkeypatch.setenv("INITIAL_ADMIN_USERNAME", "admin")
    monkeypatch.setenv("INITIAL_ADMIN_PASSWORD", "password123")
    monkeypatch.setenv("INITIAL_ADMIN_DISPLAY_NAME", "测试管理员")
    init_db()
    ensure_initial_admin()
    return db_path


@pytest.fixture
def client():
    from index import app

    with TestClient(app) as test_client:
        login_resp = test_client.post(
            "/api/auth/login",
            json={"username": "admin", "password": "password123"},
        )
        assert login_resp.status_code == 200
        yield test_client


def build_requirement_docx() -> bytes:
    document = Document()
    document.add_paragraph("4.1 功能描述")
    document.add_paragraph("新增投保资格校验，未满足资格条件时禁止提交，并提示失败原因。")
    document.add_paragraph("4.4 界面")
    document.add_paragraph("资格校验失败时，页面需要展示明显提示文案和弹窗引导。")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def create_project(client: TestClient, name: str) -> dict:
    response = client.post("/api/projects", json={"name": name, "description": ""})
    assert response.status_code == 200
    return response.json()["data"]


def create_local_user(client: TestClient, username: str, display_name: str = "项目成员") -> dict:
    response = client.post(
        "/api/users",
        json={
            "username": username,
            "password": "Reader123!",
            "display_name": display_name,
            "role": "user",
        },
    )
    assert response.status_code == 200
    return response.json()


def seed_config_assets(project_id: int, label: str) -> tuple[dict, dict, bytes]:
    requirement_content = (
        f"# {label}需求文档\n"
        f"## 4.1 功能描述\n"
        f"- {label}项目的资格校验与提交流程。\n"
    ).encode("utf-8")
    requirement_document = upsert_requirement_document(
        content_hash=hashlib.sha256(f"doc:{label}:{project_id}".encode("utf-8")).hexdigest(),
        file_name=f"{label}-requirement.md",
        file_type="md",
        file_size=len(requirement_content),
        content=requirement_content,
        source_page="权限测试",
        project_id=project_id,
        operator_username="admin",
        operator_display_name="测试管理员",
    )

    cases = [
        {
            "case_id": f"TC-{project_id:03d}",
            "description": f"{label}项目资格校验",
            "steps": "1. 输入测试数据\n2. 点击提交",
            "expected_result": "系统按预期返回结果",
            "source": "seed",
        }
    ]
    test_case_asset = upsert_test_case_asset(
        content_hash=hashlib.sha256(f"asset:{label}:{project_id}".encode("utf-8")).hexdigest(),
        asset_type="generated",
        name=f"{label}测试用例",
        file_type="json",
        file_size=0,
        cases=cases,
        source_page="权限测试",
        project_id=project_id,
        operator_username="admin",
        operator_display_name="测试管理员",
    )
    return requirement_document, test_case_asset, requirement_content


def test_case_generation_assets_are_deduplicated_only_after_save(client: TestClient):
    requirement_bytes = build_requirement_docx()
    project = create_project(client, "案例生成项目")

    map_response = client.post(
        "/api/functional-testing/case-generation/map",
        data={"project_id": str(project["id"])},
        files={
            "requirement_file": (
                "requirement.docx",
                requirement_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert map_response.status_code == 200
    mapping_snapshot = map_response.json()["data"]

    with patch(
        "services.requirement_case_generator.call_deepseek",
        new=AsyncMock(
            return_value={
                "result": {
                    "summary": "覆盖资格校验失败场景。",
                    "cases": [
                        {
                            "case_id": "TC-001",
                            "description": "资格校验失败时禁止提交",
                            "steps": "1. 输入不满足条件的数据\n2. 点击提交",
                            "expected_result": "系统阻止提交，并提示资格校验失败原因。",
                        }
                    ],
                },
                "usage": {
                    "prompt_tokens": 80,
                    "completion_tokens": 40,
                    "total_tokens": 120,
                    "prompt_cache_hit_tokens": 0,
                    "prompt_cache_miss_tokens": 80,
                },
                "provider": "DeepSeek",
                "provider_key": "deepseek",
            }
        ),
    ):
        preview_response = client.post(
            "/api/functional-testing/case-generation/generate",
            data={
                "project_id": str(project["id"]),
                "prompt_template_key": "requirement",
                "mapping_result_snapshot": json.dumps(mapping_snapshot, ensure_ascii=False),
            },
            files={
                "requirement_file": (
                    "requirement.docx",
                    requirement_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    assert preview_response.status_code == 200
    generation_snapshot = preview_response.json()["data"]
    assert generation_snapshot["total"] == 1

    for _ in range(2):
        save_response = client.post(
            "/api/functional-testing/case-generation/save",
            data={
                "project_id": str(project["id"]),
                "prompt_template_key": "requirement",
                "case_name": "资格校验回归包",
                "iteration_version": "2026.04",
                "mapping_result_snapshot": json.dumps(mapping_snapshot, ensure_ascii=False),
                "generation_result_snapshot": json.dumps(generation_snapshot, ensure_ascii=False),
                "source_page": "案例生成",
            },
            files={
                "requirement_file": (
                    "requirement.docx",
                    requirement_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
        assert save_response.status_code == 200

    documents_response = client.get("/api/config-management/requirement-documents")
    assert documents_response.status_code == 200
    documents = documents_response.json()["data"]
    assert len(documents) == 1
    assert documents[0]["file_name"] == "requirement.docx"
    assert documents[0]["project_name"] == "案例生成项目"
    assert documents[0]["source_page"] == "案例生成"
    assert documents[0]["operator_username"] == "admin"

    assets_response = client.get("/api/config-management/test-cases")
    assert assets_response.status_code == 200
    assets = assets_response.json()["data"]
    assert len(assets) == 1
    assert assets[0]["asset_type"] == "generated"
    assert assets[0]["project_name"] == "案例生成项目"
    assert assets[0]["case_count"] == 1
    assert assets[0]["source_page"] == "案例生成"
    assert assets[0]["operator_username"] == "admin"
    assert assets[0]["name"] == "资格校验回归包"
    assert assets[0]["iteration_version"] == "2026.04"

    detail_response = client.get(f"/api/config-management/test-cases/{assets[0]['id']}")
    assert detail_response.status_code == 200
    detail = detail_response.json()["data"]
    assert detail["cases"][0]["case_id"] == "TC-001"
    assert detail["cases"][0]["description"] == "资格校验失败时禁止提交"
    assert detail["iteration_version"] == "2026.04"


def test_case_quality_assets_are_deduplicated(client: TestClient):
    create_project_response = client.post("/api/projects", json={"name": "案例质检项目", "description": ""})
    assert create_project_response.status_code == 200
    project_id = create_project_response.json()["data"]["id"]

    requirement_response = client.post(
        "/api/requirement-analysis/analyze",
        data={
            "project_id": str(project_id),
            "use_ai": "false",
            "source_page": "案例质检",
        },
        files={
            "requirement_file": (
                "quality-requirement.docx",
                build_requirement_docx(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert requirement_response.status_code == 200

    code_file = FIXTURES_DIR / "sample_code_changes.json"
    test_file = FIXTURES_DIR / "sample_test_cases.csv"
    mapping_file = FIXTURES_DIR / "sample_mapping.csv"

    for _ in range(2):
        with open(code_file, "rb") as cf, open(test_file, "rb") as tf, open(mapping_file, "rb") as mf:
            analyze_response = client.post(
                f"/api/projects/{project_id}/analyze",
                data={"use_ai": "false", "source_page": "案例质检"},
                files={
                    "code_changes": ("code.json", cf, "application/json"),
                    "test_cases_file": ("tests.csv", tf, "text/csv"),
                    "mapping_file": ("mapping.csv", mf, "text/csv"),
                },
            )
        assert analyze_response.status_code == 200

    documents_response = client.get("/api/config-management/requirement-documents")
    assert documents_response.status_code == 200
    documents = documents_response.json()["data"]
    assert len(documents) == 1
    assert documents[0]["file_name"] == "quality-requirement.docx"
    assert documents[0]["project_name"] == "案例质检项目"
    assert documents[0]["source_page"] == "案例质检"
    assert documents[0]["operator_username"] == "admin"

    assets_response = client.get("/api/config-management/test-cases")
    assert assets_response.status_code == 200
    assets = assets_response.json()["data"]
    assert len(assets) == 1
    assert assets[0]["asset_type"] == "upload"
    assert assets[0]["project_name"] == "案例质检项目"
    assert assets[0]["source_page"] == "案例质检"
    assert assets[0]["operator_username"] == "admin"
    assert assets[0]["case_count"] == 4
    assert assets[0]["iteration_version"] is None

    detail_response = client.get(f"/api/config-management/test-cases/{assets[0]['id']}")
    assert detail_response.status_code == 200
    detail = detail_response.json()["data"]
    assert detail["cases"][0]["case_id"]
    assert detail["cases"][0]["description"]
    assert detail["iteration_version"] is None


def test_requirement_documents_support_original_file_download(client: TestClient):
    requirement_bytes = build_requirement_docx()
    project = create_project(client, "需求下载项目")

    analyze_response = client.post(
        "/api/requirement-analysis/analyze",
        data={
            "project_id": str(project["id"]),
            "use_ai": "false",
            "source_page": "独立需求分析",
        },
        files={
            "requirement_file": (
                "需求说明.docx",
                requirement_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert analyze_response.status_code == 200

    documents_response = client.get("/api/config-management/requirement-documents")
    assert documents_response.status_code == 200
    documents = documents_response.json()["data"]
    assert len(documents) == 1

    download_response = client.get(
        f"/api/config-management/requirement-documents/{documents[0]['id']}/download"
    )
    assert download_response.status_code == 200
    assert download_response.content == requirement_bytes
    assert (
        download_response.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "attachment;" in download_response.headers["content-disposition"]
    assert "filename*" in download_response.headers["content-disposition"]
    assert ".docx" in download_response.headers["content-disposition"]


def test_requirement_documents_support_markdown_download(client: TestClient):
    requirement_bytes = (
        "# 需求说明\n"
        "## 4.1 功能描述\n"
        "- 新增资格校验，未满足条件时禁止提交。\n"
        "## 4.4 界面\n"
        "- 页面显示资格校验失败提示。\n"
    ).encode("utf-8")
    project = create_project(client, "Markdown需求项目")

    analyze_response = client.post(
        "/api/requirement-analysis/analyze",
        data={
            "project_id": str(project["id"]),
            "use_ai": "false",
            "source_page": "案例质检",
        },
        files={
            "requirement_file": (
                "需求说明.md",
                requirement_bytes,
                "text/markdown",
            )
        },
    )
    assert analyze_response.status_code == 200

    documents_response = client.get("/api/config-management/requirement-documents")
    assert documents_response.status_code == 200
    documents = documents_response.json()["data"]
    assert len(documents) == 1

    download_response = client.get(
        f"/api/config-management/requirement-documents/{documents[0]['id']}/download"
    )
    assert download_response.status_code == 200
    assert download_response.content == requirement_bytes
    assert download_response.headers["content-type"] == "text/markdown; charset=utf-8"
    assert ".md" in download_response.headers["content-disposition"]


def test_config_management_assets_are_filtered_by_project_membership(client: TestClient):
    member = create_local_user(client, "config-reader")
    assigned_project = create_project_record(name="可见项目", tester_ids=[member["id"]])
    hidden_project = create_project_record(name="不可见项目")

    visible_document, visible_asset, visible_content = seed_config_assets(assigned_project["id"], "可见")
    hidden_document, hidden_asset, _ = seed_config_assets(hidden_project["id"], "隐藏")

    logout_response = client.post("/api/auth/logout")
    assert logout_response.status_code == 200

    login_response = client.post(
        "/api/auth/login",
        json={"username": "config-reader", "password": "Reader123!"},
    )
    assert login_response.status_code == 200

    documents_response = client.get("/api/config-management/requirement-documents")
    assets_response = client.get("/api/config-management/test-cases")
    visible_download_response = client.get(
        f"/api/config-management/requirement-documents/{visible_document['id']}/download"
    )
    hidden_download_response = client.get(
        f"/api/config-management/requirement-documents/{hidden_document['id']}/download"
    )
    visible_asset_response = client.get(f"/api/config-management/test-cases/{visible_asset['id']}")
    hidden_asset_response = client.get(f"/api/config-management/test-cases/{hidden_asset['id']}")

    assert documents_response.status_code == 200
    assert [item["id"] for item in documents_response.json()["data"]] == [visible_document["id"]]
    assert documents_response.json()["data"][0]["project_name"] == "可见项目"

    assert assets_response.status_code == 200
    assert [item["id"] for item in assets_response.json()["data"]] == [visible_asset["id"]]
    assert assets_response.json()["data"][0]["project_name"] == "可见项目"

    assert visible_download_response.status_code == 200
    assert visible_download_response.content == visible_content
    assert hidden_download_response.status_code == 404
    assert hidden_download_response.json()["detail"] == "需求文档不存在"

    assert visible_asset_response.status_code == 200
    assert visible_asset_response.json()["data"]["id"] == visible_asset["id"]
    assert hidden_asset_response.status_code == 404
    assert hidden_asset_response.json()["detail"] == "测试用例记录不存在"
