from __future__ import annotations

import copy
import io
import json
import re
from typing import Any
from urllib.parse import urlparse
from zipfile import BadZipFile

import yaml
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from services.api_automation_prompt import build_document_parse_messages
from services.deepseek_client import call_deepseek
from services.file_parser import detect_word_content_type

try:
    import fitz
except ImportError:
    fitz = None

try:
    from services.requirement_document_parser import _extract_doc_text
except Exception:
    _extract_doc_text = None


HTTP_METHODS = {"GET", "POST", "PUT", "DELETE", "PATCH", "HEAD", "OPTIONS"}
FULL_URL_RE = re.compile(r"https?://[^\s\"'<>]+", re.IGNORECASE)
PATH_RE = re.compile(r"/[A-Za-z0-9_\-./{}]+")
METHOD_RE = re.compile(r"\b(GET|POST|PUT|DELETE|PATCH|HEAD|OPTIONS)\b", re.IGNORECASE)
PURE_INDEX_RE = re.compile(r"^[\d.]+$")
TEXT_DOCUMENT_AI_TIMEOUT_SECONDS = 100
JSON_CONTEXT_WINDOW = 32
SECTION_CONTEXT_LOOKBACK_LINES = 4
REQUEST_JSON_HINTS = ("request", "body", "payload", "params", "入参", "请求", "参数")
RESPONSE_JSON_HINTS = ("response", "result", "响应", "返回", "出参")
GENERIC_ENDPOINT_NAME_PREFIXES = (
    "接口地址",
    "接口路径",
    "请求地址",
    "请求路径",
    "路径",
    "地址",
    "url",
    "uri",
    "api地址",
)


def _infer_type(value: Any) -> str:
    if isinstance(value, bool):
        return "boolean"
    if isinstance(value, int):
        return "integer"
    if isinstance(value, float):
        return "number"
    if isinstance(value, list):
        return "array"
    if isinstance(value, dict):
        return "object"
    return "string"


def _parameter_from_value(name: str, value: Any, location: str = "body", required: bool = False) -> dict[str, Any]:
    return {
        "name": name,
        "type": _infer_type(value),
        "required": required,
        "description": "",
        "example": value,
        "location": location,
    }


def _extract_parameter_constraints(schema: dict[str, Any], parameter: dict[str, Any]) -> dict[str, Any]:
    constraints: dict[str, Any] = {}
    for key in ("format", "pattern", "default", "minimum", "maximum"):
        value = schema.get(key)
        if value is not None:
            constraints[key] = value

    enum_values = schema.get("enum")
    if isinstance(enum_values, list) and enum_values:
        constraints["enum"] = enum_values

    example = parameter.get("example")
    if example is None:
        example = schema.get("example")
    if example is not None:
        constraints["example"] = example

    length_mapping = {
        "minLength": "min_length",
        "maxLength": "max_length",
        "minItems": "min_items",
        "maxItems": "max_items",
    }
    for source_key, target_key in length_mapping.items():
        value = schema.get(source_key)
        if value is not None:
            constraints[target_key] = value
    return constraints


def _schema_from_value(value: Any) -> Any:
    if isinstance(value, dict):
        return {key: _schema_from_value(item) for key, item in value.items()}
    if isinstance(value, list):
        return [] if not value else [_schema_from_value(value[0])]
    return {"type": _infer_type(value), "example": value}


def _extract_path_from_url(full_url: str) -> str:
    parsed = urlparse(full_url)
    path = (parsed.path or "").strip()
    return "" if path in {"", "/"} else path


def _looks_like_endpoint_declaration_prefix(prefix: str) -> bool:
    candidate = prefix.strip().lower()
    if not candidate:
        return True
    if candidate.endswith((":","：")):
        return True
    if any(token in candidate for token in ("get", "post", "put", "delete", "patch", "head", "options", "path", "url", "uri", "api")):
        return True
    return len(candidate) <= 8 and all(token not in candidate for token in ('"', "{", "["))


def _extract_paths_from_text_line(line: str) -> list[str]:
    stripped = str(line or "").strip()
    if not stripped:
        return []

    candidates: list[str] = []
    seen: set[str] = set()

    for url_match in FULL_URL_RE.finditer(stripped):
        full_url = url_match.group(0).rstrip("，。；;,")
        path = _extract_path_from_url(full_url)
        if _is_probable_endpoint_path(path) and path not in seen:
            seen.add(path)
            candidates.append(path)

    if candidates:
        return candidates

    for path_match in PATH_RE.finditer(stripped):
        path = path_match.group(0).rstrip("，。；;,")
        if not _is_probable_endpoint_path(path):
            continue

        prefix = stripped[:path_match.start()]
        if not _looks_like_endpoint_declaration_prefix(prefix):
            continue
        if any(token in prefix for token in ('"', "{", "[")) and not prefix.strip().endswith((":","：")):
            continue

        if path not in seen:
            seen.add(path)
            candidates.append(path)

    return candidates


def _normalize_endpoint_path_key(path: str) -> str:
    candidate = str(path or "").strip()
    if len(candidate) > 1:
        candidate = candidate.rstrip("/")
    return candidate


def _is_probable_endpoint_path(path: str) -> bool:
    candidate = str(path or "").strip()
    if not candidate or candidate == "/" or not candidate.startswith("/"):
        return False

    normalized = candidate.lstrip("/")
    if not normalized:
        return False

    first_segment = normalized.split("/", 1)[0]
    if "." in first_segment and "/" not in normalized:
        return False
    return True


def _is_meaningful_endpoint_name_candidate(value: str) -> bool:
    candidate = re.sub(r"\s+", " ", str(value or "")).strip("：:|- ")
    if not candidate:
        return False
    if PURE_INDEX_RE.fullmatch(candidate):
        return False
    if FULL_URL_RE.fullmatch(candidate):
        return False
    if any(
        candidate.startswith(prefix)
        for prefix in ("请求方式", "请求方法", "基础地址", "地址", "接口路径", "路径", "请求示例", "返回示例")
    ):
        return False
    return True


def _is_generic_endpoint_label(value: str, path: str = "") -> bool:
    candidate = re.sub(r"\s+", "", str(value or "")).strip("：:|- ").lower()
    if not candidate:
        return True

    normalized_path = _normalize_endpoint_path_key(path).strip().lower()
    normalized_path_without_slash = normalized_path.lstrip("/")
    if candidate in {normalized_path, normalized_path_without_slash}:
        return True

    return any(candidate.startswith(prefix) for prefix in GENERIC_ENDPOINT_NAME_PREFIXES)


def _resolve_text_endpoint_name(prev_lines: list[str], filename: str, endpoint_index: int, fallback_path: str) -> str:
    explicit_name = next(
        (
            line.split("：", 1)[-1].split(":", 1)[-1].strip()
            for line in reversed(prev_lines)
            if any(keyword in line for keyword in ("接口名称", "接口名", "名称"))
            and _is_meaningful_endpoint_name_candidate(line.split("：", 1)[-1].split(":", 1)[-1].strip())
        ),
        "",
    )
    if explicit_name:
        return explicit_name

    interface_name = next(
        (line for line in reversed(prev_lines) if "接口" in line and _is_meaningful_endpoint_name_candidate(line)),
        "",
    )
    if interface_name:
        return interface_name

    fallback_name = next(
        (line for line in reversed(prev_lines[-3:]) if len(line) <= 40 and _is_meaningful_endpoint_name_candidate(line)),
        "",
    )
    if fallback_name:
        return fallback_name

    return fallback_path or f"{filename} 接口{endpoint_index}"


def _normalize_endpoint(endpoint: dict[str, Any], index: int, source_type: str) -> dict[str, Any]:
    method = str(endpoint.get("method") or "POST").upper()
    if method not in HTTP_METHODS:
        method = "POST"
    raw_path = str(endpoint.get("path") or "").strip()
    path = _extract_path_from_url(raw_path) if raw_path.lower().startswith(("http://", "https://")) else raw_path
    missing_fields = list(dict.fromkeys(endpoint.get("missing_fields") or []))
    name = str(endpoint.get("name") or endpoint.get("summary") or path or f"接口{index}").strip()
    return {
        "endpoint_id": str(endpoint.get("endpoint_id") or f"endpoint-{index:03d}"),
        "group_name": str(endpoint.get("group_name") or "默认分组"),
        "name": name,
        "method": method,
        "path": path,
        "summary": str(endpoint.get("summary") or name),
        "headers": endpoint.get("headers") or [],
        "path_params": endpoint.get("path_params") or [],
        "query_params": endpoint.get("query_params") or [],
        "body_schema": endpoint.get("body_schema") or {},
        "response_schema": endpoint.get("response_schema") or {},
        "error_codes": endpoint.get("error_codes") or [],
        "dependency_hints": endpoint.get("dependency_hints") or [],
        "missing_fields": missing_fields,
        "source_type": source_type,
    }


def _endpoint_quality_score(endpoint: dict[str, Any]) -> int:
    score = 0
    if endpoint.get("path"):
        score += 4
    if not _is_generic_endpoint_label(str(endpoint.get("name") or ""), str(endpoint.get("path") or "")):
        score += 6
    if not _is_generic_endpoint_label(str(endpoint.get("summary") or ""), str(endpoint.get("path") or "")):
        score += 2
    if endpoint.get("method") in {"POST", "PUT", "PATCH"}:
        score += 1
    if endpoint.get("body_schema"):
        score += 4
        if endpoint.get("method") in {"POST", "PUT", "PATCH"}:
            score += 2
    if endpoint.get("response_schema"):
        score += 4
    score += min(len(endpoint.get("headers") or []), 3)
    score += min(len(endpoint.get("query_params") or []), 3)
    score += min(len(endpoint.get("path_params") or []), 3)
    score += min(len(endpoint.get("dependency_hints") or []), 3)
    score += min(len(endpoint.get("error_codes") or []), 3)
    score -= len(endpoint.get("missing_fields") or []) * 2
    return score


def _recalculate_missing_fields(endpoint: dict[str, Any], candidates: list[dict[str, Any]]) -> list[str]:
    merged_fields = list(
        dict.fromkeys(
            field
            for item in candidates
            for field in (item.get("missing_fields") or [])
        )
    )
    resolved_fields: set[str] = set()
    if endpoint.get("query_params") or endpoint.get("body_schema"):
        resolved_fields.add("request_params")
    if endpoint.get("body_schema"):
        resolved_fields.add("body_schema")
    if endpoint.get("response_schema"):
        resolved_fields.add("response_schema")
    return [field for field in merged_fields if field not in resolved_fields]


def _merge_text_endpoint_group(candidates: list[dict[str, Any]]) -> dict[str, Any]:
    ranked_candidates = sorted(candidates, key=_endpoint_quality_score, reverse=True)
    merged = copy.deepcopy(ranked_candidates[0])

    for candidate in ranked_candidates[1:]:
        if _is_generic_endpoint_label(merged.get("name") or "", merged.get("path") or "") and not _is_generic_endpoint_label(candidate.get("name") or "", candidate.get("path") or ""):
            merged["name"] = candidate.get("name") or merged.get("name")
        if _is_generic_endpoint_label(merged.get("summary") or "", merged.get("path") or "") and not _is_generic_endpoint_label(candidate.get("summary") or "", candidate.get("path") or ""):
            merged["summary"] = candidate.get("summary") or merged.get("summary")
        if not merged.get("body_schema") and candidate.get("body_schema"):
            merged["body_schema"] = copy.deepcopy(candidate.get("body_schema"))
        if not merged.get("response_schema") and candidate.get("response_schema"):
            merged["response_schema"] = copy.deepcopy(candidate.get("response_schema"))
        if not merged.get("headers") and candidate.get("headers"):
            merged["headers"] = copy.deepcopy(candidate.get("headers"))
        if not merged.get("query_params") and candidate.get("query_params"):
            merged["query_params"] = copy.deepcopy(candidate.get("query_params"))
        if not merged.get("path_params") and candidate.get("path_params"):
            merged["path_params"] = copy.deepcopy(candidate.get("path_params"))
        if not merged.get("error_codes") and candidate.get("error_codes"):
            merged["error_codes"] = copy.deepcopy(candidate.get("error_codes"))
        if candidate.get("method") in {"POST", "PUT", "PATCH"} and candidate.get("body_schema"):
            merged["method"] = candidate.get("method")

        merged["dependency_hints"] = list(
            dict.fromkeys([*(merged.get("dependency_hints") or []), *(candidate.get("dependency_hints") or [])])
        )

    merged["missing_fields"] = _recalculate_missing_fields(merged, ranked_candidates)
    return merged


def _dedupe_text_endpoints_by_path(endpoints: list[dict[str, Any]]) -> list[dict[str, Any]]:
    grouped: dict[str, list[dict[str, Any]]] = {}
    ordered_keys: list[str] = []

    for endpoint in endpoints:
        path_key = _normalize_endpoint_path_key(str(endpoint.get("path") or ""))
        key = path_key or str(endpoint.get("endpoint_id") or "")
        if key not in grouped:
            grouped[key] = []
            ordered_keys.append(key)
        grouped[key].append(endpoint)

    deduped: list[dict[str, Any]] = []
    for key in ordered_keys:
        candidates = grouped[key]
        if len(candidates) == 1:
            deduped.append(candidates[0])
            continue
        deduped.append(_merge_text_endpoint_group(candidates))
    return deduped


def _extract_docx_text(content: bytes) -> str:
    try:
        document = Document(io.BytesIO(content))
    except BadZipFile as exc:
        raise ValueError(
            "当前文件不是有效的 .docx 文档，请确认文件未损坏，且不是仅修改扩展名后的旧版 Word 文档。"
        ) from exc
    except PackageNotFoundError as exc:
        raise ValueError("当前 Word 文档无法打开，请确认上传的是标准 .docx 文件。") from exc

    chunks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                chunks.append(" | ".join(cells))
    return "\n".join(chunks)


def _extract_word_text(content: bytes, filename: str) -> str:
    content_type = detect_word_content_type(content)
    if content_type == "doc":
        if _extract_doc_text is None:
            raise ValueError("当前环境不支持旧版 .doc 文档解析")
        return _extract_doc_text(content)
    if content_type == "docx":
        return _extract_docx_text(content)

    if filename.lower().endswith(".docx"):
        return _extract_docx_text(content)
    if _extract_doc_text is None:
        raise ValueError("当前环境不支持旧版 .doc 文档解析")
    return _extract_doc_text(content)


def _extract_pdf_text(content: bytes) -> str:
    if fitz is None:
        raise RuntimeError("PDF 解析依赖 PyMuPDF 未安装")
    document = fitz.open(stream=content, filetype="pdf")
    try:
        chunks = [page.get_text().strip() for page in document]
    finally:
        document.close()
    text = "\n".join(chunk for chunk in chunks if chunk)
    if not text.strip():
        raise ValueError("扫描件/不可提取文本，当前不支持")
    return text


def _parse_openapi_document(spec: dict[str, Any], source_type: str) -> list[dict[str, Any]]:
    paths = spec.get("paths") or {}
    endpoints: list[dict[str, Any]] = []
    global_security = spec.get("security") or []
    for path, path_item in paths.items():
        if not isinstance(path_item, dict):
            continue
        for method, operation in path_item.items():
            if str(method).upper() not in HTTP_METHODS or not isinstance(operation, dict):
                continue

            parameters = operation.get("parameters") or []
            query_params: list[dict[str, Any]] = []
            path_params: list[dict[str, Any]] = []
            headers: list[dict[str, Any]] = []
            for parameter in parameters:
                if not isinstance(parameter, dict):
                    continue
                location = str(parameter.get("in") or "")
                schema = parameter.get("schema") or {}
                item = {
                    "name": str(parameter.get("name") or ""),
                    "type": (schema.get("type") or "string"),
                    "required": bool(parameter.get("required", False)),
                    "description": str(parameter.get("description") or ""),
                    "location": location,
                }
                item.update(_extract_parameter_constraints(schema, parameter))
                if location == "query":
                    query_params.append(item)
                elif location == "path":
                    path_params.append(item)
                elif location == "header":
                    headers.append(item)

            body_schema: dict[str, Any] = {}
            request_body = operation.get("requestBody") or {}
            content_section = request_body.get("content") or {}
            json_body = content_section.get("application/json") or next(iter(content_section.values()), {})
            if isinstance(json_body, dict):
                schema = json_body.get("schema")
                if isinstance(schema, dict):
                    body_schema = schema

            response_schema: dict[str, Any] = {}
            responses = operation.get("responses") or {}
            error_codes: list[dict[str, Any]] = []
            for status_code, response in responses.items():
                if not isinstance(response, dict):
                    continue
                content_map = response.get("content") or {}
                json_response = content_map.get("application/json") or next(iter(content_map.values()), {})
                if not response_schema and isinstance(json_response, dict):
                    schema = json_response.get("schema")
                    if isinstance(schema, dict):
                        response_schema = schema
                if str(status_code).startswith(("4", "5")) or str(status_code) == "default":
                    error_codes.append({"code": str(status_code), "description": str(response.get("description") or "")})

            dependency_hints = ["需要鉴权"] if operation.get("security") or global_security else []
            missing_fields = []
            if not body_schema and str(method).upper() in {"POST", "PUT", "PATCH"}:
                missing_fields.append("body_schema")
            if not response_schema:
                missing_fields.append("response_schema")

            endpoints.append(_normalize_endpoint({
                "endpoint_id": f"{method.lower()}-{path.strip('/').replace('/', '-') or 'root'}",
                "group_name": (
                    (operation.get("tags") or ["默认分组"])[0]
                    if isinstance(operation.get("tags"), list) and operation.get("tags")
                    else "默认分组"
                ),
                "name": operation.get("summary") or operation.get("operationId") or path,
                "method": method,
                "path": path,
                "summary": operation.get("description") or operation.get("summary") or path,
                "headers": headers,
                "path_params": path_params,
                "query_params": query_params,
                "body_schema": body_schema,
                "response_schema": response_schema,
                "error_codes": error_codes,
                "dependency_hints": dependency_hints,
                "missing_fields": missing_fields,
            }, len(endpoints) + 1, source_type))
    return endpoints


def _find_json_block_with_bounds(text: str, marker_index: int = 0) -> tuple[Any, int, int] | None:
    search_index = marker_index
    while True:
        start = text.find("{", search_index)
        if start == -1:
            return None

        depth = 0
        for index in range(start, len(text)):
            char = text[index]
            if char == "{":
                depth += 1
            elif char == "}":
                depth -= 1
                if depth == 0:
                    block = text[start:index + 1]
                    try:
                        return json.loads(block), start, index + 1
                    except json.JSONDecodeError:
                        break
        search_index = start + 1


def _find_json_block(text: str, marker_index: int = 0) -> tuple[Any, int] | None:
    match = _find_json_block_with_bounds(text, marker_index)
    if match is None:
        return None
    value, _, end = match
    return value, end


def _collect_json_blocks(text: str) -> list[tuple[Any, int, int]]:
    blocks: list[tuple[Any, int, int]] = []
    search_index = 0
    while True:
        match = _find_json_block_with_bounds(text, search_index)
        if match is None:
            return blocks
        blocks.append(match)
        _, _, end = match
        search_index = end


def _extract_request_response_examples(section_text: str) -> tuple[Any, Any]:
    blocks = _collect_json_blocks(section_text)
    if not blocks:
        return {}, {}

    request_json: Any = {}
    response_json: Any = {}
    unlabeled_blocks: list[Any] = []

    for block, start, _ in blocks:
        prefix = section_text[max(0, start - JSON_CONTEXT_WINDOW):start].lower()
        if any(keyword in prefix for keyword in RESPONSE_JSON_HINTS):
            if not response_json:
                response_json = block
            continue
        if any(keyword in prefix for keyword in REQUEST_JSON_HINTS):
            if not request_json:
                request_json = block
            continue
        unlabeled_blocks.append(block)

    if not request_json and unlabeled_blocks:
        request_json = unlabeled_blocks.pop(0)
    if not response_json and unlabeled_blocks:
        response_json = unlabeled_blocks.pop(0)
    return request_json, response_json


def _build_text_endpoint(raw_text: str, filename: str) -> list[dict[str, Any]]:
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    if not lines:
        return []

    anchors: list[tuple[int, str]] = []
    for line_index, line in enumerate(lines):
        for path in _extract_paths_from_text_line(line):
            anchors.append((line_index, path))

    if anchors:
        parsed_endpoints: list[dict[str, Any]] = []
        for anchor_position, (line_index, path) in enumerate(anchors):
            next_line_index = anchors[anchor_position + 1][0] if anchor_position + 1 < len(anchors) else len(lines)
            prev_start = max(0, line_index - SECTION_CONTEXT_LOOKBACK_LINES)
            prev_lines = lines[prev_start:line_index]
            context_lines = lines[prev_start:next_line_index]
            section_lines = lines[line_index:next_line_index]
            context_text = "\n".join(context_lines)
            section_text = "\n".join(section_lines)
            request_json, response_json = _extract_request_response_examples(section_text)
            method_match = METHOD_RE.search(section_text) or METHOD_RE.search(context_text)
            method = method_match.group(1).upper() if method_match else ("POST" if request_json else "GET")
            endpoint_index = len(parsed_endpoints) + 1
            endpoint_name = _resolve_text_endpoint_name(prev_lines, filename, endpoint_index, path)

            missing_fields = []
            if not request_json:
                missing_fields.append("request_params")
            if not response_json:
                missing_fields.append("response_schema")

            dependency_hints = []
            lower_text = context_text.lower()
            if "token" in lower_text or "cookie" in lower_text or "sign" in lower_text or "签名" in context_text:
                dependency_hints.append("需要鉴权或签名")

            parsed_endpoints.append(_normalize_endpoint({
                "endpoint_id": f"endpoint-{endpoint_index:03d}",
                "group_name": "文档解析",
                "name": endpoint_name,
                "method": method,
                "path": path,
                "summary": endpoint_name,
                "headers": [],
                "path_params": [],
                "query_params": [],
                "body_schema": _schema_from_value(request_json) if isinstance(request_json, (dict, list)) else {},
                "response_schema": _schema_from_value(response_json) if isinstance(response_json, (dict, list)) else {},
                "error_codes": [],
                "dependency_hints": dependency_hints,
                "missing_fields": missing_fields,
            }, endpoint_index, "text_document"))

        return parsed_endpoints

    endpoints: list[dict[str, Any]] = []
    search_index = 0
    endpoint_index = 1
    while True:
        url_match = FULL_URL_RE.search(raw_text, search_index)
        if not url_match:
            break
        full_url = url_match.group(0).rstrip("，。；;,")
        path = _extract_path_from_url(full_url)
        if not _is_probable_endpoint_path(path):
            search_index = url_match.end()
            continue
        nearby_text = raw_text[max(0, url_match.start() - 500): min(len(raw_text), url_match.end() + 3000)]
        method_match = METHOD_RE.search(nearby_text)
        method = method_match.group(1).upper() if method_match else ("POST" if "{" in nearby_text else "GET")
        request_json_match = _find_json_block(raw_text, url_match.end())
        request_json = request_json_match[0] if request_json_match else {}
        next_index = request_json_match[1] if request_json_match else url_match.end()
        response_json_match = _find_json_block(raw_text, next_index)
        response_json = response_json_match[0] if response_json_match else {}

        prev_text_window = raw_text[max(0, url_match.start() - 240): url_match.start()]
        prev_lines = [line for line in lines if line in prev_text_window]
        endpoint_name = _resolve_text_endpoint_name(prev_lines, filename, endpoint_index, path)

        missing_fields = []
        if not request_json:
            missing_fields.append("request_params")
        if not response_json:
            missing_fields.append("response_schema")

        dependency_hints = []
        lower_text = nearby_text.lower()
        if "token" in lower_text or "cookie" in lower_text or "sign" in lower_text or "签名" in nearby_text:
            dependency_hints.append("需要鉴权或签名")

        endpoints.append(_normalize_endpoint({
            "endpoint_id": f"endpoint-{endpoint_index:03d}",
            "group_name": "文档解析",
            "name": endpoint_name,
            "method": method,
            "path": path,
            "summary": endpoint_name,
            "headers": [],
            "path_params": [],
            "query_params": [],
            "body_schema": _schema_from_value(request_json) if isinstance(request_json, dict) else {},
            "response_schema": _schema_from_value(response_json) if isinstance(response_json, dict) else {},
            "error_codes": [],
            "dependency_hints": dependency_hints,
            "missing_fields": missing_fields,
        }, endpoint_index, "text_document"))
        endpoint_index += 1
        search_index = url_match.end()

    if endpoints:
        return endpoints

    for path_match in PATH_RE.finditer(raw_text):
        path = path_match.group(0).rstrip("，。；;,")
        if not _is_probable_endpoint_path(path):
            continue
        request_json_match = _find_json_block(raw_text, path_match.end())
        request_json = request_json_match[0] if request_json_match else {}
        prev_text_window = raw_text[max(0, path_match.start() - 240): path_match.start()]
        prev_lines = [line for line in lines if line in prev_text_window]
        endpoint_name = _resolve_text_endpoint_name(prev_lines, filename, 1, path)
        return [_normalize_endpoint({
            "endpoint_id": "endpoint-001",
            "group_name": "文档解析",
            "name": endpoint_name,
            "method": "POST" if request_json else "GET",
            "path": path,
            "summary": endpoint_name,
            "headers": [],
            "path_params": [],
            "query_params": [],
            "body_schema": _schema_from_value(request_json) if isinstance(request_json, dict) else {},
            "response_schema": {},
            "error_codes": [],
            "dependency_hints": [],
            "missing_fields": ["response_schema"],
        }, 1, "text_document")]
    return []


async def _enhance_text_endpoints_with_ai(
    filename: str,
    raw_text: str,
    endpoints: list[dict[str, Any]],
    prompt_template_text: str | None = None,
) -> list[dict[str, Any]]:
    ai_response = await call_deepseek(
        build_document_parse_messages(filename, raw_text, prompt_template_text=prompt_template_text),
        timeout_seconds=TEXT_DOCUMENT_AI_TIMEOUT_SECONDS,
        max_retries=0,
    )
    ai_result = ai_response.get("result") or {}
    ai_endpoints = ai_result.get("endpoints") or []
    if not isinstance(ai_endpoints, list):
        return _dedupe_text_endpoints_by_path(endpoints)
    normalized_ai = [
        _normalize_endpoint(item, index + 1, "text_document_ai")
        for index, item in enumerate(ai_endpoints)
        if isinstance(item, dict)
    ]
    normalized_ai = [item for item in normalized_ai if _is_probable_endpoint_path(item.get("path") or "")]
    if not normalized_ai:
        return _dedupe_text_endpoints_by_path(endpoints)

    return _dedupe_text_endpoints_by_path([*endpoints, *normalized_ai])


async def parse_api_document(
    content: bytes,
    filename: str,
    use_ai: bool = True,
    prompt_template_text: str | None = None,
) -> dict[str, Any]:
    lower_name = filename.lower()
    source_type = "openapi" if lower_name.endswith((".json", ".yaml", ".yml")) else "text_document"

    if lower_name.endswith(".pdf"):
        raw_text = _extract_pdf_text(content)
        endpoints = _build_text_endpoint(raw_text, filename)
    elif lower_name.endswith((".doc", ".docx")):
        raw_text = _extract_word_text(content, filename)
        endpoints = _build_text_endpoint(raw_text, filename)
    elif lower_name.endswith(".json"):
        payload = json.loads(content.decode("utf-8"))
        raw_text = json.dumps(payload, ensure_ascii=False, indent=2)
        endpoints = _parse_openapi_document(payload, "openapi_json")
    elif lower_name.endswith((".yaml", ".yml")):
        payload = yaml.safe_load(content.decode("utf-8"))
        raw_text = yaml.safe_dump(payload, allow_unicode=True, sort_keys=False)
        endpoints = _parse_openapi_document(payload or {}, "openapi_yaml")
    else:
        raise ValueError("当前仅支持 PDF / Word / OpenAPI(JSON/YAML) 接口文档")

    if use_ai and raw_text and source_type == "text_document":
        try:
            endpoints = await _enhance_text_endpoints_with_ai(
                filename,
                raw_text,
                endpoints,
                prompt_template_text=prompt_template_text,
            )
        except Exception:
            pass

    if source_type == "text_document":
        endpoints = _dedupe_text_endpoints_by_path(endpoints)

    if not endpoints:
        raise ValueError("未能从文档中识别出接口，请检查文档格式或补充更完整的接口描述")

    missing_fields = sorted({field for item in endpoints for field in item.get("missing_fields", [])})
    return {
        "file_name": filename,
        "source_type": source_type,
        "raw_text": raw_text,
        "raw_text_excerpt": raw_text[:1200],
        "endpoint_count": len(endpoints),
        "missing_fields": missing_fields,
        "endpoints": endpoints,
    }
